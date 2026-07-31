from openai import OpenAI
import random
import time
import textwrap
import os
import pandas as pd
from docx import Document
from fpdf import FPDF
from services.mocktest.PromptsDict import prompt_templates
from datetime import datetime
import json
import uuid
# import winsound
from pymongo import MongoClient, errors
from pymongo.errors import ConnectionFailure, OperationFailure
from typing import List
import certifi
import cloudinary
import cloudinary.uploader

# ===============================================================
# === ROBUST PATH CONFIGURATION ===
# ===============================================================

# Get the absolute path of the directory containing this script (Generation.py)
# This will be .../src/MockTestAutomation/
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Get the project root directory by going up two levels from the script's directory
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, '..', '..'))

# Build absolute paths to ensure files are always found
# MODEL = 'gpt-4-turbo'
MODEL = os.getenv("OPENAI_MODEL", "gpt-4-turbo")
# SAVE_GENERATIONS_TO_DB = True
SAVE_GENERATIONS_TO_DB = True
API_KEY = os.getenv("OPENAI_API_KEY") or os.getenv("API_KEY")
# questions_per_chunk = 5 
# questions_per_chunk = 3 # change to 5 when on production level

# EXCEL_PATH = os.path.join(SCRIPT_DIR, "..", "data", "Syllabus.xlsx")
# EXCEL_PATH = os.path.join(SCRIPT_DIR, "..", "data", "UGCSyllabus.xlsx")

# Define the path to the backend/data directory
BACKEND_DATA_DIR = os.getenv(
    "BACKEND_DATA_DIR", 
    "/app/data" if os.path.exists("/app") else os.path.join(PROJECT_ROOT, "data")
)
# Create the output directories inside backend/data
OUTPUT_DIR = os.path.join(BACKEND_DATA_DIR, "generated_files")
RAW_RESPONSES_DIR = os.path.join(BACKEND_DATA_DIR, "raw_responses")

# Ensure the output directories exist at the project root
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(RAW_RESPONSES_DIR, exist_ok=True)

# ===============================================================
# === MONGODB SETUP ===
# ===============================================================
try:
    MONGO_CONNECTION_STRING = os.getenv("MONGO_URI")
    if not MONGO_CONNECTION_STRING:
        print("⚠️ MONGO_URI not set. Logging disabled.")
        mongo_client = None
    else:
        # Added tlsAllowInvalidCertificates for local dev SSL issues
        mongo_client = MongoClient(
            MONGO_CONNECTION_STRING,
            serverSelectionTimeoutMS=5000,
            tls=True,
            tlsCAFile=certifi.where(), # Tells the container where the SSL certs are
            tlsAllowInvalidCertificates=True,
            connectTimeoutMS=10000 
        )
        # Verify connection
        mongo_client.admin.command('ping')
        print("✅ MongoDB connection successful.")
        db = mongo_client["acetrack_finetune_db"]
        finetune_collection = db["generation_logs"]
except Exception as e:
    print(f"❌ MongoDB connection failed: {type(e).__name__} - {e}")
    mongo_client = None
# ===============================================================
if os.getenv("CLOUDINARY_URL"):
    cloudinary.config(cloudinary_url=os.getenv("CLOUDINARY_URL"))
else:
    cloudinary.config(
        cloud_name=os.getenv("CLOUDINARY_CLOUD_NAME"),
        api_key=os.getenv("CLOUDINARY_API_KEY"),
        api_secret=os.getenv("CLOUDINARY_API_SECRET"),
        secure=True
    )
# === TOPIC LOADING AND PROMPT GENERATION ===
# def load_all_topics(excel_path=EXCEL_PATH):
#     """Loads and shuffles topics from the specified Excel file."""
#     if not os.path.exists(excel_path):
#         raise FileNotFoundError(f"Syllabus file not found at the expected path: {excel_path}. Please ensure it exists.")
#     try:
#         df = pd.read_excel(excel_path)
#         if df.empty or len(df.columns) == 0:
#             raise ValueError("Syllabus.xlsx is empty or has no columns.")
#         topic_column = df.columns[0]
#         topics = df[topic_column].dropna().tolist()
#         random.shuffle(topics)
#         return topics
#     except Exception as e:
#         # Catch other potential pandas errors
#         raise RuntimeError(f"Failed to read or process the Excel file at {excel_path}: {e}")

def validate_topic_capacity(plan, topics_dict: dict, questions_per_chunk: int):
    """Validates if there are enough topics for the requested questions."""
    total_topics_count = sum(len(subj_topics) for subj_topics in topics_dict.values())
    total_chunks_requested = sum(num // questions_per_chunk for num in plan.values())
    if total_chunks_requested > total_topics_count // questions_per_chunk:
        error_msg = f"Not enough unique topics to generate the requested number of questions. \nTopics available: {total_topics_count}, Questions requested: {sum(plan.values())}"
        raise ValueError(error_msg)

def build_prompt_from_template(topics_list, template_key, num_of_questions, EXAM):
    """Builds a GPT prompt from a template with the given topics."""
    topics_str = "\n".join([f"{i+1}. {topic}" for i, topic in enumerate(topics_list)])
    randomized_answer_key = ', '.join(str(n) for n in random.choices(range(1, 5), k=num_of_questions))
    template = prompt_templates.get(template_key, "")
    return template.format(topics=topics_str, answer_key=randomized_answer_key, num=num_of_questions, exam=EXAM)

def generate_all_prompts(plan, topics: dict, exam, questions_per_chunk: int):
    """Generates a list of all prompts to be sent to the GPT API."""
    prompts = []
    
    # 1. Shuffle topics within each subject
    shuffled_by_subj = {}
    for subj, subj_topics in topics.items():
        shuffled_by_subj[subj] = random.sample(subj_topics, len(subj_topics))
        
    # 2. Interleave them perfectly for maximum diversity
    shuffled_topics = []
    subjects = list(shuffled_by_subj.keys())
    
    while any(shuffled_by_subj.values()):
        for subj in subjects:
            if shuffled_by_subj[subj]:
                shuffled_topics.append(shuffled_by_subj[subj].pop(0))
                
    topic_index = 0
    
    for qtype, count in plan.items():
        if topic_index + ((count // questions_per_chunk) * questions_per_chunk) > len(shuffled_topics):
            raise ValueError("Topic index out of bounds. This indicates a logic error in topic validation.")
        for _ in range(count // questions_per_chunk):
            chunk = shuffled_topics[topic_index : topic_index + questions_per_chunk]
            topic_index += questions_per_chunk
            prompt = build_prompt_from_template(chunk, qtype, questions_per_chunk, exam)
            prompts.append((qtype, prompt))
    return prompts

# === FILE OPERATIONS ===
def save_to_docx(content, filename):
    """Saves the given content to a .docx file in the output directory."""
    path = os.path.join(OUTPUT_DIR, filename)
    try:
        doc = Document()
        doc.add_paragraph(content)
        doc.save(path)
    except Exception as e:
        raise IOError(f"❌ Cannot access {path}. Details: {e}")
    
def save_to_pdf(content, filename):
    """Saves the given content to a .pdf file in the output directory."""
    path = os.path.join(OUTPUT_DIR, filename)
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        encoded_content = content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 5, txt=encoded_content)
        pdf.output(path)
    except Exception as e:
        raise IOError(f"❌ Cannot save PDF to {path}. Details: {e}")

def save_raw_response(text, folder=RAW_RESPONSES_DIR):
    """Saves the raw GPT response for debugging purposes."""
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    file_id = uuid.uuid4().hex[:6]
    # filename = f"gpt_response_{timestamp}.docx"
    filename = f"gpt_response_{timestamp}_{file_id}.pdf"
    path = os.path.join(folder, filename)
    # doc = Document()
    # doc.add_paragraph(text)
    # doc.save(path)
    # print(f"✅ Raw response saved to: {path}")
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        encoded_text = text.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 5, txt=encoded_text)
        pdf.output(path)
        print(f"✅ Raw response saved to: {path}")
    except Exception as e:
        print(f"❌ Failed to save raw response PDF. Details: {e}")

def log_generation_to_db(system_prompt: str, user_prompt: str, response_content: str, exam_name: str, model_name: str, testing: bool):
    """Logs a successful prompt/response pair to MongoDB if enabled."""
    if not testing and mongo_client and SAVE_GENERATIONS_TO_DB and response_content:
        try:
            finetune_collection.insert_one({
                "system": system_prompt,
                "prompt": user_prompt,
                "response": response_content,
                "exam": exam_name,
                "model": model_name,
                "created_at": datetime.now()
            })
            print("  -> ✅ Logged generation pair to MongoDB.")
        except OperationFailure as e:
            print(f"  -> ⚠️ Failed to log to MongoDB (OperationFailure): {e.details}")
        except Exception as e:
            print(f"  -> ⚠️ Failed to log to MongoDB (General Error): {e}")
    elif not SAVE_GENERATIONS_TO_DB:
        print("  -> 🟡 Skipping MongoDB log (master flag is OFF).")

# === GPT HANDLING ===
def call_gpt(prompt, testing, exam_name, chunks, retries=3):
    """Calls the OpenAI API with a given prompt, with retries."""
    system_prompt = f"You are a {exam_name} paper setter."
    
    if testing:
        time.sleep(1)
        mock_response = "\n\n".join([
            f"--Question Starting--\nQ{i+1}. This is a sample test question for type {prompt[:3]}.\nAnswer: A\nExplanation: This is a test explanation."
            for i in range(chunks)
        ])
        return mock_response, system_prompt
    
    api_key = os.getenv("OPENAI_API_KEY") or os.getenv("API_KEY") or API_KEY
    if not api_key:
        raise ValueError("OPENAI_API_KEY environment variable not set.")

    client = OpenAI(api_key=api_key)
    for attempt in range(retries):
        try:
            response = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=4000
            )
            response_content = response.choices[0].message.content

            return response_content, system_prompt
        
        except Exception as e:
            print(f"⚠️ GPT attempt {attempt+1} failed: {e}")
            if attempt < retries - 1:
                time.sleep(2)
    raise RuntimeError("❌ All GPT API retries failed.")


# === CORE EXECUTION LOGIC ===
def handle_generation(prompts, TESTING, exam_name, questions_per_chunk: int):
    """Handles the question generation loop, calling GPT for each prompt."""
    all_questions = []
    skipped_chunks = []
    # questions_per_chunk = 5
    max_retries_per_chunk = 3
    
    for qtype, prompt in prompts:
        print(f"Generating questions for type: {qtype}")
        generated_chunk = None
        last_failed_chunk = []
        
        for attempt in range(max_retries_per_chunk):
            try:
                print(f"  -> Attempt {attempt + 1} for {qtype}...")
                response, system_prompt_used = call_gpt(prompt, TESTING, exam_name, questions_per_chunk)
                
                if not response:
                    print(f"  -> ⚠️ call_gpt returned empty response for {qtype}.")
                    last_failed_chunk = [f"--- GPT CALL RETURNED EMPTY RESPONSE --- (Type: {qtype})"]
                    continue
                
                if not TESTING:
                    save_raw_response(response)

                # Parse questions by delimiter
                if "--Question Starting--" in response:
                    questions = [q.strip() for q in textwrap.dedent(response).split("--Question Starting--") if q.strip()]
                else:
                    # Fallback: split by double newlines if delimiter wasn't used
                    questions = [q.strip() for q in textwrap.dedent(response).split("\n\n") if q.strip()]

                last_failed_chunk = questions if questions else [response]
                
                # --- VALIDATION LOGIC ---
                if len(questions) >= questions_per_chunk:
                    print(f"  ✅ Success! Got {len(questions)} questions.")
                    generated_chunk = questions[:questions_per_chunk]
                    if not TESTING and system_prompt_used:
                         log_generation_to_db(
                             system_prompt=system_prompt_used,
                             user_prompt=prompt,
                             response_content=response,
                             exam_name=exam_name,
                             model_name=MODEL,
                             testing=TESTING
                         )
                    break
                else:
                    print(f"  ⚠️ Validation warning: GPT returned {len(questions)} questions instead of {questions_per_chunk}.")
                    if attempt == max_retries_per_chunk - 1 and len(questions) > 0:
                        # Accept whatever we got on the final attempt
                        print(f"  ⚠️ Accepting {len(questions)} questions on final retry.")
                        generated_chunk = questions
                        break
                    time.sleep(1)

            except Exception as e:
                print(f"An error occurred during GPT call for {qtype}: {e}")
                last_failed_chunk = [f"--- ERROR IN GENERATION ---: {str(e)}"]
                if attempt < max_retries_per_chunk - 1:
                    time.sleep(2)
        
        # After the retry loop, check if we got a valid chunk
        if generated_chunk:
            all_questions.extend(generated_chunk)
        else:
            print(f"❌ Failed to generate a valid chunk for {qtype} after {max_retries_per_chunk} attempts. Skipping.")
            skipped_chunks.append(last_failed_chunk)
            
    return all_questions, skipped_chunks


# === MAIN ENTRY POINT FOR BACKEND ===
def run_generation_task(plan: dict, testing_mode: bool, exam_name: str, output_format: str, questions_per_chunk: int, topics: dict):
    """Main function to be called by the FastAPI """
    try:
        print(f"Starting generation for {exam_name} with plan: {plan}")
        
        run_id = uuid.uuid4().hex[:8]
        # questions_filename = f"Questions_{run_id}.docx"
        # skipped_filename = f"Skipped_{run_id}.docx"
        extension = ".pdf" if output_format == 'pdf' else ".docx"
        questions_filename = f"Questions_{run_id}{extension}"
        skipped_filename = f"Skipped_{run_id}{extension}"
        
        save_function = save_to_pdf if output_format == 'pdf' else save_to_docx

        # topics = load_all_topics()
        
        validate_topic_capacity(plan, topics, questions_per_chunk)
        
        prompts = generate_all_prompts(plan, topics, exam_name, questions_per_chunk)
        
        generated_questions, skipped_chunks = handle_generation(prompts, testing_mode, exam_name, questions_per_chunk)
        if not generated_questions and not skipped_chunks:
            raise RuntimeError("No questions were successfully generated. Check logs for API errors or response format issues.")
        
        # save_to_docx("\n\n".join(generated_questions), questions_filename)
        # save_to_pdf("\n\n".join(generated_questions), questions_filename)
        generated_files = {}
        message = ""
        if generated_questions:
            save_function("\n\n".join(generated_questions), questions_filename)
            # --- CLOUDINARY UPLOAD: Questions ---
            message = f"Successfully generated {len(generated_questions)} questions."
            try:
                print(f"Uploading {questions_filename} to Cloudinary...")
                upload_questions = cloudinary.uploader.upload(
                    os.path.join(OUTPUT_DIR, questions_filename),
                    resource_type="raw",
                    public_id=f"ace-track/{questions_filename}"
                )
                # Store the Cloudinary URL instead of the local filename
                generated_files["questions"] = upload_questions.get("secure_url")
            except Exception as u_err:
                print(f"⚠️ Cloudinary upload failed for questions: {u_err}")
                generated_files["questions"] = questions_filename # Fallback to filename
            
        if skipped_chunks:
            skipped_text = "\n\n".join([
                f"--- Skipped Chunk {i+1} ---:\n" + "\n\n".join(chunk)
                for i, chunk in enumerate(skipped_chunks)
            ])
            # save_to_docx(skipped_text, skipped_filename)
            save_function(skipped_text, skipped_filename)
            # --- CLOUDINARY UPLOAD: Skipped ---
            try:
                upload_skipped = cloudinary.uploader.upload(
                    os.path.join(OUTPUT_DIR, skipped_filename),
                    resource_type="raw",
                    public_id=f"ace-track/{skipped_filename}"
                )
                generated_files["skipped"] = upload_skipped.get("secure_url")
            except Exception as u_err:
                print(f"⚠️ Cloudinary upload failed for skipped chunks: {u_err}")
                generated_files["skipped"] = skipped_filename # Fallback
            
            
        if message and len(skipped_chunks)>0: # Add to existing message
                message += f" Failed to generate {len(skipped_chunks)} chunk(s), which have been saved separately."
        elif message:
            pass
        else: # Create new message
                message = f"Failed to generate questions, but {len(skipped_chunks)} skipped chunk(s) were saved."

        print("\n✅ Mock Test Generation Completed.")
        
        # if os.name == 'nt': # 'nt' is Windows
        #     try:
        #         import winsound
        #         winsound.PlaySound("CorrectHarp.wav", winsound.SND_FILENAME)
        #     except Exception as e:
        #         print(f"🟡 Windows notification sound failed: {e}")
        # else:
        #     # In Linux/Container environments, we just log completion
        #     print("🔔 Generation Task Finished")
        
        # generated_files = {"questions": questions_filename}
        # message = "Questions generated successfully."
        # if skipped_chunks:
        #     generated_files["skipped"] = skipped_filename

        return {
            "success": True,
            "message": message,
            "files": generated_files
            # "partial_success": bool(skipped_chunks and generated_questions) # True only if we have both
        }

    except Exception as e:
        error_message = f"Question generation failed: {str(e)}"
        print(f"❌ {error_message}")
        return {"success": False, "message": error_message, "files": {}}